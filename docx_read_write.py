import docx
d = docx.Document('C:\\Users\\16616\\Documents\\demo.docx')

# prints the paragraph objects as python ogjects
print(d.paragraphs)

# prints the text of the paragraph
print(d.paragraphs[0].text)

# assigning first paragraph object to p
p = d.paragraphs[1]

# a run is a group of text with a class (bold/italic/normal/underlined)

# this will print the first set of normal text in that paragraph before font changes
print(p.runs[3].text)

# modify the document with some basic commands here
p.runs[3].italic = True
p.runs[3].underline = True
p.runs[3].text = 'underlines and italic'

# change style to Title
p.style = 'Title'

# save this as a new document

d.save('C:\\Users\\16616\\Documents\\demo3.docx')

# creates a document object within python
d2 = docx.Document()

# adding paragraphs
d2.add_paragraph('Hello this is a paragraph')
d2.add_paragraph('Hello this is another paragraph')

d2.save('C:\\Users\\16616\\Documents\\demo4.docx')

# adding a run and making the text bold
p2 = d2.paragraphs[0]
p2.add_run('This is a new run.')
p2.runs[1].bold = True

d2.save('C:\\Users\\16616\\Documents\\demo6.docx')